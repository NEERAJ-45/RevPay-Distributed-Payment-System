"""
Generate .docx files from RevPay markdown docs.
Replaces ASCII diagrams with diagram prompts.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def is_diagram_block(lines):
    """Check if a code block is an ASCII diagram vs actual code."""
    if not lines:
        return False
    diagram_chars = set('│─└├┐┌┘┤┬┴┼►▲▼→←↔↓↑╔╗╚╝║═┏┓┛┗┳┻┣┫┃╋╂╆╇╈╊┅┉┄┈')
    code_lines = [l for l in lines if l.strip() and not l.strip().startswith(('```', '~~~'))]
    if not code_lines:
        return False
    diag_count = sum(1 for l in code_lines if any(c in l for c in diagram_chars))
    # Also detect box-drawing / arrow heavy blocks
    arrow_count = sum(1 for l in code_lines if '──' in l or '►' in l or '───' in l)
    return (diag_count / len(code_lines)) > 0.3 or arrow_count > 2


def make_diagram_prompt(title, block_text):
    """Create a descriptive prompt from a diagram block."""
    lines = [l for l in block_text.split('\n') if l.strip()]
    desc_lines = [l for l in lines if not l.strip().startswith(('│', '├', '└', '┌', '─', '►', '▲')) and len(l.strip()) > 3]
    key_terms = []
    for l in lines:
        for t in ['POST', 'GET', 'PUT', 'HTTP', 'Kafka', 'DB', 'Redis', 'JWT', 'Feign', 'SMS', '→', '←']:
            if t in l:
                key_terms.append(t)
    desc = '. '.join(desc_lines[:3]) if desc_lines else ''
    return (
        f"[DIAGRAM: {title}]\n"
        f"Prompt: Create a {title.lower()} showing the data flow between components.\n"
        f"Include these elements: {', '.join(sorted(set(key_terms), key=lambda x: -len(x))[:8]) if key_terms else 'system components'}.\n"
    )


def extract_diagram_title(text):
    """Create a short title from the surrounding context."""
    lines = text.strip().split('\n')
    for l in lines:
        l = l.strip()
        if l.startswith('#'):
            return l.lstrip('#').strip()
    return "System Flow"


def convert_md_to_docx(md_text, output_path, doc_title):
    """Convert markdown to docx, handling diagram replacement."""
    doc = Document()
    
    # Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h_style.font.name = 'Calibri'
    
    # Title
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process lines
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buffer = []
    code_lang = ""
    in_table = False
    table_buffer = []
    in_list = False
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:].strip()
                code_buffer = []
            else:
                in_code = False
                block_text = '\n'.join(code_buffer)
                if is_diagram_block(code_buffer):
                    # This is an ASCII diagram - replace with prompt
                    title = extract_diagram_title('\n'.join(lines[max(0,i-20):i]))
                    prompt = make_diagram_prompt(title, block_text)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(prompt)
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.left_indent = Inches(0.3)
                else:
                    # Regular code block
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(block_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_buffer = []
            i += 1
            continue
        
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        
        stripped = line.strip()
        
        # Skip thematic breaks
        if stripped in ('---', '***', '___'):
            doc.add_paragraph('').paragraph_format.space_before = Pt(2)
            i += 1
            continue
        
        # Empty line
        if not stripped:
            i += 1
            continue
        
        # Headings
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            text = stripped.lstrip('#').strip()
            if level <= 4:
                doc.add_heading(text, level)
            else:
                doc.add_heading(text, 4)
            i += 1
            continue
        
        # Table rows
        if '|' in stripped and stripped.startswith('|'):
            parts = [p.strip() for p in stripped.split('|')[1:-1]]
            # Check if it's a separator row
            if all(all(c in '-:' for c in p) for p in parts if p):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_buffer = [parts]
            else:
                table_buffer.append(parts)
            # Check if next line is also a table
            if i + 1 < len(lines) and '|' in lines[i+1] and lines[i+1].strip().startswith('|'):
                i += 1
                continue
            # End of table - render it
            if table_buffer:
                num_cols = len(table_buffer[0]) if table_buffer else 2
                table = doc.add_table(rows=len(table_buffer), cols=num_cols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, row_data in enumerate(table_buffer):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            # Bold the first row
                            if r_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                table_buffer = []
            in_table = False
            i += 1
            continue
        
        # Unordered list
        if stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('+ '):
            text = stripped[2:]
            level = 0
            raw = line
            indent = len(line) - len(line.lstrip())
            if indent >= 4:
                level = 1
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            i += 1
            continue
        
        # Numbered list
        match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            i += 1
            continue
        
        # Regular paragraph (handle inline formatting)
        p = doc.add_paragraph()
        
        # Simple inline formatting
        bold_parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in bold_parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Check for inline code
                code_parts = re.split(r'(`[^`]+`)', part)
                for cp in code_parts:
                    if cp.startswith('`') and cp.endswith('`'):
                        r = p.add_run(cp[1:-1])
                        r.font.name = 'Consolas'
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x99, 0x2E, 0x4E)
                    else:
                        p.add_run(cp)
        
        i += 1
    
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    files = [
        ('docs/INTERVIEW_DEEP_DIVE.md', 'RevPay — Interview-Ready Deep Dive'),
        ('docs/CODEBASE_STUDY_GUIDE.md', 'RevPay Codebase Study Guide'),
        ('MICROSERVICE_FLOW.md', 'RevPay Microservice Flow — Detailed Breakdown'),
    ]
    
    out_dir = os.path.join(DOCS_DIR, 'docs')
    if not os.path.exists(out_dir):
        out_dir = DOCS_DIR
    
    for rel_path, title in files:
        src = os.path.join(DOCS_DIR, rel_path)
        base = os.path.splitext(os.path.basename(src))[0]
        dst = os.path.join(out_dir, f'{base}.docx')
        
        if os.path.exists(src):
            with open(src, 'r', encoding='utf-8') as f:
                md_text = f.read()
            convert_md_to_docx(md_text, dst, title)
            print(f'Created: {dst}')
        else:
            print(f'Skipped (not found): {src}')
