import docx

DOC_PATH = r'E:\Study\Main-Content\SigmaWebDev\SigmaWebDev\PROJECTS\RevPay---Distributed-Payment-System\docs\INTERVIEW_DEEP_DIVE.docx'
OUT_PATH = r'E:\Study\Main-Content\SigmaWebDev\SigmaWebDev\PROJECTS\RevPay---Distributed-Payment-System\docs\INTERVIEW_DEEP_DIVE_UPDATED.docx'

prompts = [
    # 0 - User Registration Happy Path
    "Generate an Excalidraw diagram JSON for the RevPay User Registration flow.\n"
    "Layout: Client -> User Service -> PostgreSQL -> Kafka\n"
    "Flow:\n"
    "1. Client sends POST /api/auth/register {fullName, phone, pin} to User Service\n"
    "2. User Service validates input (phone pattern, pin 4-6 digits)\n"
    "3. User Service checks uniqueness: SELECT by phone -> PostgreSQL\n"
    "4. PostgreSQL returns (nil) or existing user\n"
    "5. User Service calls UpiIdGenerator: sanitize(fullName).toLowerCase() + \"@miniupi\"\n"
    "6. User Service hashes PIN with BCryptPasswordEncoder(12)\n"
    "7. User Service saves user: INSERT INTO users -> PostgreSQL (OK)\n"
    "8. User Service saves outbox event: INSERT INTO outbox_events (processed=false) -> PostgreSQL\n"
    "9. User Service generates JWT (subject=userId, claims=upiId, phone, 24h expiry)\n"
    "10. User Service returns 201 + JWT + upiId to Client\n"
    "Include a right-side callout box listing the steps with JWT generation highlighted.",

    # 1 - 1.2 Edge Cases
    "Generate an Excalidraw diagram JSON with 3 rows for 3 failure scenarios:\n\n"
    "Row 1: Scenario A - Duplicate Phone Number\n"
    "Client -> User Service: POST /api/auth/register {phone}\n"
    "User Service -> PostgreSQL: SELECT by phone\n"
    "PostgreSQL -> User Service: returns existing User\n"
    "User Service -> Client: 409 CONFLICT {error: PHONE_ALREADY_EXISTS}\n\n"
    "Row 2: Scenario D - Outbox Write Fails\n"
    "Step 1: INSERT INTO users -> OK\n"
    "Step 2: INSERT INTO outbox_events -> FAIL\n"
    "Dashed red arrow: @Transactional ROLLBACK over Step 1\n\n"
    "Row 3: Scenario E - Kafka Down\n"
    "OutboxScheduler --x Kafka: Connection refused\n"
    "Circular retry arrow: every 2s\n"
    "Later: Kafka recovers -> Wallet/Notification consume events.",

    # 2 - Scenario B: Duplicate UPI ID
    "Generate an Excalidraw diagram JSON with two parallel solutions:\n\n"
    "SOLUTION 1 (current): UpiIdGenerator checks DB, appends counter -> \"alice2@miniupi\"\n"
    "SOLUTION 2 (better): Append random suffix -> \"alice_x7k2@miniupi\"\n\n"
    "Both: UNIQUE constraint in DB is final safety net.\n"
    "Arrows: UpiIdGenerator -> PostgreSQL SELECT -> INSERT with fallback.",

    # 3 - Scenario D detail
    "Generate an Excalidraw diagram JSON for outbox write failure scenario.\n\n"
    "Timeline (top-to-bottom):\n"
    "Step 1: User Service -> PostgreSQL: INSERT INTO users -> OK (green checkmark)\n"
    "Step 2: User Service -> PostgreSQL: INSERT INTO outbox_events -> FAIL (red X, constraint)\n"
    "Dashed red arrow labeled \"@Transactional ROLLBACK\" from Step 2 back over Step 1\n\n"
    "Callout box:\n"
    "Result: Step 1 ROLLS BACK - user NOT created\n"
    "Client receives 500 error\n"
    "Client retries -> idempotent (phone unique -> 409, not duplicate)\n"
    "\"Either both succeed, or neither. Never one without the other.\"",

    # 4 - Scenario E detail
    "Generate an Excalidraw diagram JSON with 3 phases:\n\n"
    "PHASE 1: User Service -> PostgreSQL: INSERT user + outbox_event (processed=false) - green check\n\n"
    "PHASE 2: OutboxScheduler --x Kafka: publishSync() - Connection refused (red dashed)\n"
    "Circular retry arrow: \"Retry every 2s\"\n\n"
    "PHASE 3: Kafka recovers (green)\n"
    "OutboxScheduler -> Kafka: publishSync() - ACK (green dashed)\n"
    "Kafka -> Wallet Service: user.created -> createWallet(upiId)\n"
    "Kafka -> Notification Service: user.created -> sendWelcome()\n"
    "Notification -> User Phone: [SMS] Welcome to RevPay!\n\n"
    "Bottom: \"The outbox pattern gives at-least-once delivery even when Kafka is down.\"",

    # 5 - 1.3 Design Decisions
    "Generate an Excalidraw diagram JSON comparing two approaches side by side:\n\n"
    "LEFT: Anti-Pattern - Dual Write\n"
    "User Service -> PostgreSQL: userRepository.save(user) - OK\n"
    "User Service --x Kafka: kafkaTemplate.send() - FAILS (red dashed)\n"
    "Red label: \"DUAL-WRITE PROBLEM - User saved but no event\"\n\n"
    "RIGHT: Correct Pattern - Transactional Outbox\n"
    "User Service -> PostgreSQL: userRepository.save(user)\n"
    "User Service -> PostgreSQL: outboxEventRepository.save(outbox) - SAME TX\n"
    "Bracket: \"SAME database transaction - succeed or fail together\"\n"
    "OutboxScheduler (clock icon) - dashed green -> Kafka: publishSync()\n"
    "Kafka -> Wallet Service: createWallet()\n"
    "Kafka -> Notification Service: sendWelcome()\n"
    "Green check: \"At-least-once delivery guaranteed\"",

    # 6 - Why separate DB per service?
    "Generate an Excalidraw diagram JSON with 3 columns:\n\n"
    "Column 1: User Service -> DB: upi_users (users, outbox_events)\n"
    "Column 2: Wallet Service -> DB: upi_wallets (wallets @Version, ledger_entries)\n"
    "Column 3: Transaction Service -> DB: upi_transactions (transactions, outbox_events)\n\n"
    "Green benefits: Loose coupling, Independent scaling, Data ownership, Isolation\n"
    "Orange drawback: Cross-service queries need API calls (no JOIN across DBs)",

    # 7 - Wallet Happy Path
    "Generate an Excalidraw diagram JSON for the Wallet Add Money flow.\n\n"
    "Client -> GW -> Wallet Service: POST /wallet/add-money/alice@miniupi {amount: 5000, note: Salary}\n"
    "Wallet Service: findByUpiId() -> returns wallet (balance=0, version=1)\n"
    "setBalance(0 + 5000) = 5000\n"
    "Wallet Service -> PostgreSQL: UPDATE wallets SET balance=5000, version=2 WHERE id=uuid AND version=1\n"
    "Wallet Service: create LedgerEntry CREDIT amount=5000 balanceAfter=5000\n"
    "Wallet Service -> PostgreSQL: INSERT ledger_entries\n"
    "Wallet Service -> Client: 200 OK {balance: 5000}\n\n"
    "Callout: Ledger entries are append-only - immutable audit trail.",

    # 8 - Concurrent Add-Money
    "Generate an Excalidraw diagram JSON with Thread A and Thread B side by side:\n\n"
    "THREAD A (left):\n"
    "read balance=0 (version=1) -> set 5000 -> UPDATE wallets SET balance=5000, version=2 WHERE id=uuid AND version=1 -> SUCCESS (green)\n\n"
    "THREAD B (right):\n"
    "read balance=0 (version=1) -> set 3000 -> UPDATE wallets SET balance=3000, version=2 WHERE id=uuid AND version=1 -> FAILS (0 rows affected, red)\n"
    "Dashed orange arrow: @Retryable -> retry -> read balance=5000 (version=2) -> set 8000 -> SUCCESS\n\n"
    "Bottom comparison:\n"
    "Without @Version: Thread B OVERWRITES -> Rs 5000 lost\n"
    "With @Version: OptimisticLockException -> retry -> correct -> money safe",

    # 9 - Insufficient Funds
    "Generate an Excalidraw diagram JSON for insufficient funds scenario.\n\n"
    "Client -> Transaction Service: POST /transactions/pay {amount: 200}\n"
    "Transaction Service -> Feign -> Wallet Service: POST /wallet/internal/transfer {amount: 200}\n"
    "Wallet Service: balance=100 < amount=200 -> throws InsufficientFundsException (red arrow)\n"
    "Transaction Service: @Transactional ROLLBACK\n"
    "Transaction Service -> PostgreSQL: UPDATE status=FAILED, failure_reason=Insufficient funds\n"
    "Transaction Service -> PostgreSQL: INSERT outbox_event (TRANSACTION_FAILED)\n"
    "OutboxScheduler -> Kafka: publish txn.failed\n"
    "Kafka -> Notification Service -> User Phone: [SMS] Payment failed. Insufficient funds.\n\n"
    "Callout: sender.getBalance().compareTo(amount) < 0 -> @Transactional rolls back everything",

    # 10 - Atomic Transfer Happy Path
    "Generate an Excalidraw diagram JSON for the Atomic Transfer.\n\n"
    "@Transactional boundary box containing 6 steps:\n"
    "Step 1: findByUpiId(fromUpiId) -> locks sender (with @Version)\n"
    "Step 2: Check balance >= amount? YES\n"
    "Step 3: sender.setBalance(balance - amount) -> PostgreSQL: UPDATE wallets SET balance=..., version++\n"
    "Step 4: findByUpiId(toUpiId) -> locks receiver\n"
    "Step 5: receiver.setBalance(balance + amount) -> PostgreSQL: UPDATE wallets SET balance=..., version++\n"
    "Step 6: INSERT debit + credit ledger entries (append-only, immutable)\n\n"
    "Label: \"Either ALL 6 steps succeed, or ALL are rolled back\"",

    # 11 - @Transactional on DB not HTTP
    "Generate an Excalidraw diagram JSON for the transaction boundary design decision.\n\n"
    "Transaction Service -> Feign HTTP POST -> Wallet Service\n"
    "Wallet Service's @Transactional covers ONLY its own DB\n"
    "Once HTTP 200 sent, money IS moved\n\n"
    "CRASH scenario:\n"
    "Step 1: Feign call succeeds (money moved) - HTTP 200 OK\n"
    "Step 2: CRASH (red lightning bolt) - txn.setStatus(SUCCESS) never executes\n"
    "State: Money in Bob's wallet, Transaction PENDING in DB, No Kafka event\n\n"
    "3 solution boxes below:\n"
    "1. Idempotent wallet transfer (check transactionId)\n"
    "2. Saga with compensation (reverseTransfer)\n"
    "3. Distributed saga orchestrator",

    # 12 - Idempotency Check Step 1
    "Generate an Excalidraw diagram JSON for the Idempotency Check.\n\n"
    "Two lanes:\n\n"
    "LANE 1: First time (cache miss)\n"
    "Transaction Service -> Redis: GET idempotency:req-001\n"
    "Redis -> Transaction Service: (nil) - cache miss\n"
    "Transaction Service: proceed to process payment\n\n"
    "LANE 2: Replayed request (cache hit)\n"
    "Transaction Service -> Redis: GET idempotency:req-001\n"
    "Redis -> Transaction Service: cached txnId\n"
    "Transaction Service -> PostgreSQL: findById(txnId)\n"
    "PostgreSQL -> Transaction Service: existing SUCCESS txn\n"
    "Transaction Service -> Client: PayResponse {replayed: true}\n\n"
    "Callout: Safe replay - no money moved, no fraud check, no Feign call\n"
    "Bottom: Redis ~0.1ms vs Postgres ~1-5ms. TTL 24h.",

    # 13 - Save PENDING Step 2
    "Generate an Excalidraw diagram JSON for Save PENDING transaction.\n\n"
    "Transaction Service -> PostgreSQL: INSERT INTO transactions (status='PENDING')\n\n"
    "Right callout:\n"
    "Why PENDING first?\n"
    "If crash after wallet transfer but before SUCCESS:\n"
    "- PENDING record exists -> reconciliation can recover\n"
    "- Options: check wallet, mark FAILED, retry\n"
    "Without PENDING: ghost transfers - money moved but no record",

    # 14 - Mark SUCCESS Step 5
    "Generate an Excalidraw diagram JSON for Mark SUCCESS + Store Idempotency.\n\n"
    "Two parallel arrows from Transaction Service:\n"
    "Left: txn.setStatus(SUCCESS) -> PostgreSQL: UPDATE transactions WHERE id=uuid\n"
    "Right: idempotencyService.storeResult() -> Redis: SET idempotency:req-001 <txnId> EX 86400\n\n"
    "Annotation box:\n"
    "DB = source of truth (if Redis down, recover from DB)\n"
    "Redis = performance cache (100x faster, TTL auto-cleans)\n"
    "Cache-aside: 1. Check cache 2. Miss? Check DB 3. Write through",

    # 15 - Critical Question
    "Generate an Excalidraw diagram JSON for the critical distributed systems question.\n\n"
    "Timeline:\n"
    "1: Transaction Service -Feign-> Wallet Service: POST /wallet/internal/transfer\n"
    "2: Wallet Service processes -> HTTP 200 OK (money MOVED)\n"
    "3: CRASH (red lightning bolt) - status=SUCCESS never executes\n\n"
    "State after crash (red box):\n"
    "- Wallet: Alice -200, Bob +200 (money moved)\n"
    "- Transaction: PENDING in DB\n"
    "- Redis: no idempotency key\n"
    "- Kafka: no event\n\n"
    "If retry: DOUBLE CHARGE risk!\n\n"
    "3 solutions:\n"
    "1. Idempotent Wallet Transfer\n"
    "2. Saga with Compensation\n"
    "3. Distributed Transaction Coordinator",

    # 16 - Why PENDING first
    "Generate an Excalidraw diagram JSON for the PENDING state decision.\n\n"
    "Decision tree from center box \"Save as PENDING\":\n"
    "Branch A (green): Wallet Feign SUCCESS -> UPDATE status=SUCCESS -> happy path\n"
    "Branch B (orange): Feign success but crash -> PENDING record exists -> reconcile\n"
    "  - Check wallet if transfer completed\n"
    "  - Mark FAILED and refund\n"
    "  - Retry the transfer\n\n"
    "Without PENDING: crash -> no record -> ghost transfer -> money lost",

    # 17 - Why TWO writes
    "Generate an Excalidraw diagram JSON for the dual-write decision.\n\n"
    "Transaction Service with two arrows:\n"
    "Arrow 1 -> PostgreSQL: UPDATE status=SUCCESS (source of truth)\n"
    "Arrow 2 -> Redis: SET idempotency:<key>=<value> EX 86400 (performance cache)\n\n"
    "If Redis down: fallback to DB findByRequestId()\n"
    "If DB fails: whole TX rolls back\n\n"
    "Tiered: 1. Check Redis (fast) 2. Miss? Check DB (slow) 3. Write through both",

    # 18 - Write Outbox Step 6
    "Generate an Excalidraw diagram JSON for Step 6: Write Outbox Event.\n\n"
    "Transaction Service inside @Transactional boundary with saveOutboxEvent()\n"
    "Two branches:\n"
    "Branch A: INSERT outbox_events (event_type=TRANSACTION_SUCCESS, processed=false) -> PostgreSQL\n"
    "Branch B: INSERT outbox_events (event_type=TRANSACTION_FAILED, processed=false) -> PostgreSQL\n\n"
    "Boundary box:\n"
    "Step 5 + Step 6 both succeed -> committed\n"
    "Step 5 succeeds + Step 6 fails -> Step 5 ROLLS BACK\n"
    "Transaction stays PENDING, client retries -> safe",

    # 19 - fixedDelay vs fixedRate
    "Generate an Excalidraw diagram JSON comparing fixedDelay vs fixedRate.\n\n"
    "LEFT: fixedRate(2000ms) - BAD\n"
    "0s: Run 1 starts (500 events, 10s)\n"
    "2s: Run 2 OVERLAPS\n"
    "4s: Run 3 OVERLAPS\n"
    "Red warning: Concurrent processing of same events!\n\n"
    "RIGHT: fixedDelay(2000ms) - GOOD\n"
    "0s: Run 1 starts\n"
    "10s: Run 1 finishes\n"
    "12s: Run 2 starts\n"
    "Green check: No overlap, predictable load, backpressure",

    # 20 - Scheduler crash
    "Generate an Excalidraw diagram JSON for scheduler crash after publish.\n\n"
    "PHASE 1: Crash\n"
    "OutboxScheduler -> Kafka: publishSync() -> Kafka ACKs (event in Kafka)\n"
    "CRASH (lightning bolt) -> setProcessed(true) never executes\n\n"
    "PHASE 2: Recovery\n"
    "Restart -> SELECT WHERE processed=false -> same event picked AGAIN\n"
    "OutboxScheduler -> Kafka: publishes DUPLICATE\n\n"
    "PHASE 3: Consumer\n"
    "Notification: check txnId already sent? -> yes -> skip (idempotent)\n\n"
    "Bottom: Outbox = AT-LEAST-ONCE, consumer MUST handle duplicates",

    # 21 - Kafka Consumer Groups
    "Generate an Excalidraw diagram JSON for Kafka Consumer Groups.\n\n"
    "Center: Kafka with 3 topics\n"
    "- user.created (3 partitions, 7 day retention)\n"
    "- txn.completed (3 partitions)\n"
    "- txn.failed (3 partitions)\n\n"
    "Left: Wallet Service (wallet-service-group)\n"
    "Arrow: user.created -> UserCreatedListener -> createWallet()\n\n"
    "Right: Notification Service (notification-service-group)\n"
    "Arrow 1: user.created -> UserCreatedListener -> sendWelcome() SMS\n"
    "Arrow 2: txn.completed -> TransactionEventListener -> debit/credit alerts SMS\n"
    "Arrow 3: txn.failed -> TransactionEventListener -> failure alert SMS\n\n"
    "Legend: Each group reads independently, manages own offset",

    # 22 - Why separate consumer groups
    "Generate an Excalidraw diagram JSON for separate consumer groups independence.\n\n"
    "PANEL 1: Notification DOWN\n"
    "Kafka -> wallet-service-group: Wallet Service -> createWallet() -> SUCCESS\n"
    "Kafka -> notification-service-group: Notification (DOWN) -> offset NOT committed\n"
    "When back up: resumes from last offset -> processes all missed messages\n\n"
    "PANEL 2: Wallet DOWN\n"
    "Kafka -> notification-service-group: Notification -> sendWelcome() -> SUCCESS\n"
    "Kafka -> wallet-service-group: Wallet (DOWN) -> offset NOT committed\n"
    "When back up: catches up\n\n"
    "Both groups read same topic, failures isolated",

    # 23 - JWT Auth Filter
    "Generate an Excalidraw diagram JSON for JWT Auth Filter Flow.\n\n"
    "Decision tree (top-to-bottom):\n"
    "Request arrives at API Gateway\n"
    "Path in [/auth/**, /swagger-ui/**]?\n"
    "YES (green): skip JWT -> route downstream\n"
    "NO: JWT validation\n"
    "  Auth header exists?\n"
    "  NO (red): -> 401 UNAUTHORIZED\n"
    "  YES: extract Bearer token\n"
    "    Parse JWT with hmacShaKeyFor(secret)\n"
    "    Invalid/expired (red): -> 401\n"
    "    Valid (green):\n"
    "      Store upiId in exchange attributes\n"
    "      Forward to downstream\n\n"
    "Green=success paths, Red=401 rejection paths",

    # 24 - Rate Limiting
    "Generate an Excalidraw diagram JSON for Rate Limiting.\n\n"
    "Flow:\n"
    "Request -> KeyResolver (Client IP) -> Redis Token Bucket\n"
    "Bucket: burstCapacity=40, replenishRate=20/s\n"
    "Tokens > 0? YES -> consume 1 -> forward (green)\n"
    "Tokens = 0? NO -> 429 Too Many Requests (red)\n\n"
    "Table:\n"
    "/users/** | 20/s burst 40\n"
    "/wallet/** | 20/s burst 40\n"
    "/transactions/** | 10/s burst 20 (expensive payments)",

    # 25 - Q1: Double charges
    "Generate an Excalidraw diagram JSON with 3 stacked defense layers:\n\n"
    "LAYER 1: Idempotency Key (Redis)\n"
    "Client -> Transaction Service: {requestId: uuid}\n"
    "Transaction Service -> Redis: GET idempotency:uuid\n"
    "Hit -> return cached (skip processing). Miss -> proceed.\n\n"
    "LAYER 2: Unique Constraint (PostgreSQL)\n"
    "request_id has UNIQUE constraint\n"
    "If Redis down: Thread A INSERT OK, Thread B INSERT fails -> catch -> return existing\n\n"
    "LAYER 3: @Version Optimistic Locking (Wallet)\n"
    "First UPDATE WHERE version=1 -> success (version->2)\n"
    "Second UPDATE WHERE version=1 -> fails (0 rows) -> retry with new version\n\n"
    "Three independent defense layers against double charges",

    # 26 - Q2: Redis goes down
    "Generate an Excalidraw diagram JSON for Redis going down.\n\n"
    "3 impact areas:\n\n"
    "IMPACT 1: Idempotency\n"
    "Normal: Redis GET (~0.1ms)\n"
    "Redis DOWN -> fallback: PostgreSQL findByRequestId (~5ms)\n"
    "Orange fallback arrow\n\n"
    "IMPACT 2: Rate Limiting\n"
    "Normal: Redis token bucket\n"
    "Redis DOWN -> all requests pass through (no limiting)\n"
    "Red warning: Rate limiting disabled until Redis recovers\n\n"
    "IMPACT 3: JWT Auth - UNAFFECTED\n"
    "Uses local secret key, no Redis dependency\n"
    "Green check\n\n"
    "Recovery: Redis stateless, keys TTL, no permanent data loss",

    # 27 - Q3: Scaling
    "Generate an Excalidraw diagram JSON with 4 scaling panels:\n\n"
    "PANEL 1: Horizontal Scaling\n"
    "ALB -> multiple instances of each service\n"
    "Label: Each service is stateless -> scale horizontally\n\n"
    "PANEL 2: Database Sharding\n"
    "hash(sender_upi) % N -> Shard 0/1/2/N\n\n"
    "PANEL 3: Kafka Partitioning\n"
    "Topic with Partition 0/1/2 -> Consumer 0/1/2\n"
    "Label: More partitions = more parallel consumers\n\n"
    "PANEL 4: Caching\n"
    "Redis: cache daily limits (avoid SUM query), cache user profiles (avoid service calls)",
]

doc = docx.Document(DOC_PATH)

prompt_idx = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith('Prompt:') and prompt_idx < len(prompts):
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = prompts[prompt_idx]
        else:
            para.add_run(prompts[prompt_idx])
        prompt_idx += 1

doc.save(OUT_PATH)
print(f'Done. Replaced {prompt_idx} prompts. Saved to INTERVIEW_DEEP_DIVE_UPDATED.docx')
